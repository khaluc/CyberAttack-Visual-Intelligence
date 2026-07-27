"""PHASE 1 — normalize security data sources into plain UTF-8 text."""
from __future__ import annotations

import importlib.util
import json
import re
import tempfile
from dataclasses import dataclass, asdict
from email import policy
from email.parser import BytesParser
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree


TEXT_EXTENSIONS = {".txt", ".md", ".log", ".syslog", ".csv", ".json", ".xml", ".cef", ".leef"}
WORD_EXTENSIONS = {".docx", ".doc"}
PDF_EXTENSIONS = {".pdf"}
EMAIL_EXTENSIONS = {".eml", ".msg"}
WINDOWS_EVENT_EXTENSIONS = {".evtx"}
ALLOWED_EXTENSIONS = TEXT_EXTENSIONS | WORD_EXTENSIONS | PDF_EXTENSIONS | EMAIL_EXTENSIONS | WINDOWS_EVENT_EXTENSIONS


def parser_capabilities() -> dict:
    """Return truthful PHASE 1 parser availability for health checks."""
    modules = {
        "pymupdf": "fitz",
        "pdfplumber": "pdfplumber",
        "python_docx": "docx",
        "textract": "textract",
        "python_evtx": "Evtx",
    }
    engines = {
        name: importlib.util.find_spec(module) is not None
        for name, module in modules.items()
    }
    return {
        "ready": all(
            engines[name]
            for name in ("pymupdf", "pdfplumber", "python_docx", "python_evtx")
        ),
        "engines": engines,
        "extensions": sorted(ALLOWED_EXTENSIONS),
    }


@dataclass
class ParsedDocument:
    filename: str
    source_type: str
    parser: str
    text: str
    characters: int
    lines: int
    truncated: bool = False

    def to_dict(self):
        return asdict(self)


def parse_document(filename: str, raw: bytes, max_characters: int = 12000) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        supported = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise ValueError(f"Định dạng không hỗ trợ. Các định dạng hợp lệ: {supported}")

    if suffix in PDF_EXTENSIONS:
        text, parser = _parse_pdf(raw)
        source_type = "PDF"
    elif suffix in WORD_EXTENSIONS:
        text, parser = _parse_word(raw, suffix)
        source_type = "Word"
    elif suffix in EMAIL_EXTENSIONS:
        text, parser = _parse_email(raw, suffix)
        source_type = "Email"
    elif suffix in WINDOWS_EVENT_EXTENSIONS:
        text, parser = _parse_evtx(raw)
        source_type = "Windows Event"
    else:
        text, parser = _parse_security_text(raw, suffix)
        source_type = _detect_security_source(text, suffix)

    text = _normalize(text)
    if not text:
        raise ValueError("Không trích xuất được nội dung văn bản từ tệp.")
    original_length = len(text)
    text = text[:max_characters]
    return ParsedDocument(
        filename=Path(filename).name, source_type=source_type, parser=parser,
        text=text, characters=len(text), lines=text.count("\n") + 1,
        truncated=original_length > max_characters,
    )


def _parse_pdf(raw):
    errors = []
    try:
        import fitz  # PyMuPDF
        document = fitz.open(stream=raw, filetype="pdf")
        text = "\n".join(page.get_text("text") for page in document)
        if text.strip():
            return text, "PyMuPDF"
    except Exception as exc:
        errors.append(f"PyMuPDF: {exc}")
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(raw)) as document:
            text = "\n".join(page.extract_text() or "" for page in document.pages)
        if text.strip():
            return text, "pdfplumber"
    except Exception as exc:
        errors.append(f"pdfplumber: {exc}")
    text = _textract_fallback(raw, ".pdf")
    if text:
        return text, "textract"
    raise ValueError("Không đọc được PDF. " + "; ".join(errors))


def _parse_word(raw, suffix):
    if suffix == ".docx":
        try:
            from docx import Document
            document = Document(BytesIO(raw))
            blocks = [p.text for p in document.paragraphs]
            for table in document.tables:
                blocks.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            text = "\n".join(blocks)
            if text.strip():
                return text, "python-docx"
        except Exception:
            pass
    text = _textract_fallback(raw, suffix)
    if text:
        return text, "textract"
    raise ValueError("Không đọc được tài liệu Word. File .doc cũ yêu cầu textract/antiword.")


def _parse_email(raw, suffix=".eml"):
    if suffix == ".msg":
        text = _textract_fallback(raw, suffix)
        if text:
            return text, "textract"
        raise ValueError(
            "Không đọc được email Outlook .msg. Cần textract và extract-msg."
        )
    message = BytesParser(policy=policy.default).parsebytes(raw)
    parts = [
        f"From: {message.get('from', '')}", f"To: {message.get('to', '')}",
        f"Date: {message.get('date', '')}", f"Subject: {message.get('subject', '')}",
    ]
    attachments = []
    if message.is_multipart():
        for part in message.walk():
            disposition = part.get_content_disposition()
            if disposition == "attachment":
                attachments.append(part.get_filename() or "unnamed")
            elif part.get_content_type() == "text/plain":
                try:
                    parts.append(part.get_content())
                except Exception:
                    parts.append(part.get_payload(decode=True).decode(errors="replace"))
    else:
        try:
            parts.append(message.get_content())
        except Exception:
            parts.append(message.get_payload(decode=True).decode(errors="replace"))
    if attachments:
        parts.append("Attachments: " + ", ".join(attachments))
    return "\n".join(parts), "email.parser"


def _parse_evtx(raw):
    try:
        from Evtx.Evtx import Evtx
        with tempfile.NamedTemporaryFile(suffix=".evtx") as handle:
            handle.write(raw)
            handle.flush()
            with Evtx(handle.name) as log:
                events = [record.xml() for record in log.records()]
        return "\n".join(events), "python-evtx"
    except ImportError as exc:
        raise ValueError("Cần cài python-evtx để đọc Windows Event .evtx.") from exc
    except Exception as exc:
        raise ValueError(f"Windows Event không hợp lệ: {exc}") from exc


def _parse_security_text(raw, suffix):
    text = raw.decode("utf-8", errors="replace")
    if suffix == ".json":
        try:
            text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            pass
    elif suffix == ".xml":
        try:
            root = ElementTree.fromstring(text)
            text = "\n".join(f"{node.tag.split('}')[-1]}={node.text}" for node in root.iter() if node.text and node.text.strip())
        except ElementTree.ParseError:
            pass
    return text, "native text parser"


def _detect_security_source(text, suffix):
    sample = text[:6000].lower()
    if suffix == ".syslog" or re.search(r"<\d{1,3}>[a-z]{3}\s+\d+\s+\d{2}:", sample):
        return "Syslog"
    if suffix in (".cef", ".leef") or "cef:" in sample or "leef:" in sample:
        return "Firewall"
    if any(token in sample for token in ("eventid", "event id", "microsoft-windows-", "<event ")):
        return "Windows Event"
    if any(token in sample for token in ("src=", "dst=", "srcip", "dstip", "action=allow", "action=deny")):
        return "Firewall"
    return "Log" if suffix in (".log", ".csv", ".json", ".xml") else "Text"


def _textract_fallback(raw, suffix):
    try:
        import textract
    except ImportError:
        return ""
    path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
            handle.write(raw)
            path = handle.name
        return textract.process(path).decode("utf-8", errors="replace")
    except Exception:
        return ""
    finally:
        if path:
            Path(path).unlink(missing_ok=True)


def _normalize(text):
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
