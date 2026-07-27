"""Server-side incident report generators.

The structured PHASE 3 incident is the only data contract accepted here.  The
module intentionally has no Flask dependency so HTTP routes, CLI jobs and tests
can reuse the same deterministic exporters.
"""
from __future__ import annotations

import base64
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from structured_attack import validate_structured_incident


PROJECT_ROOT = Path(__file__).resolve().parent
PPTX_SCRIPT = PROJECT_ROOT / "scripts" / "generate_report_pptx.mjs"
DEFAULT_REPORT_DIR = PROJECT_ROOT / "output" / "reports"


class ReportGenerationError(RuntimeError):
    """Raised when a report runtime is missing or an exporter fails."""


def generate_pdf(
    incident: dict[str, Any],
    output_path: str | os.PathLike[str] | None = None,
    *,
    graph_image: bytes | bytearray | io.BytesIO | str | os.PathLike[str] | None = None,
    recommendations: Iterable[str] | None = None,
) -> Path:
    """Create a real server-side PDF report with ReportLab."""
    data = _canonical_incident(incident)
    destination = _prepare_output(data, output_path, ".pdf")
    graph_bytes = _read_binary(graph_image)
    actions = _recommendations(data, recommendations)

    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Image,
            KeepTogether,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ReportGenerationError(
            "PDF export requires reportlab. Install reportlab>=4."
        ) from exc

    font_regular, font_bold = _register_pdf_fonts()
    palette = {
        "navy": colors.HexColor("#0B2545"),
        "blue": colors.HexColor("#2E74B5"),
        "muted": colors.HexColor("#5E6C7B"),
        "line": colors.HexColor("#D9E1EA"),
        "surface": colors.HexColor("#F4F6F9"),
        "table": colors.HexColor("#F2F4F7"),
        "white": colors.white,
        "severity": colors.HexColor(_severity_color(data["severity"])),
    }
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CVI Body",
        parent=styles["BodyText"],
        fontName=font_regular,
        fontSize=9.5,
        leading=12,
        textColor=palette["navy"],
        spaceAfter=3,
    )
    small = ParagraphStyle(
        "CVI Small",
        parent=body,
        fontSize=7.4,
        leading=9,
        spaceAfter=0,
    )
    table_header = ParagraphStyle(
        "CVI Table Header",
        parent=small,
        fontName=font_bold,
        textColor=palette["navy"],
        alignment=TA_LEFT,
    )
    title = ParagraphStyle(
        "CVI Title",
        parent=body,
        fontName=font_bold,
        fontSize=23,
        leading=27,
        textColor=colors.black,
        spaceAfter=4,
    )
    subtitle = ParagraphStyle(
        "CVI Subtitle",
        parent=body,
        fontSize=12.5,
        leading=16,
        textColor=palette["muted"],
        spaceAfter=10,
    )
    heading = ParagraphStyle(
        "CVI Heading 1",
        parent=body,
        fontName=font_bold,
        fontSize=15,
        leading=18,
        textColor=palette["blue"],
        spaceBefore=7,
        spaceAfter=4,
        keepWithNext=True,
    )
    metric_label = ParagraphStyle(
        "CVI Metric Label",
        parent=small,
        fontName=font_bold,
        textColor=palette["muted"],
        alignment=TA_CENTER,
    )
    metric_value = ParagraphStyle(
        "CVI Metric Value",
        parent=body,
        fontName=font_bold,
        fontSize=13,
        leading=16,
        textColor=palette["navy"],
        alignment=TA_CENTER,
        spaceAfter=0,
    )

    def para(
        value: Any,
        style: ParagraphStyle = body,
        *,
        markup: bool = False,
    ) -> Paragraph:
        return Paragraph(str(value) if markup else _reportlab_text(value), style)

    document = SimpleDocTemplate(
        str(destination),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.72 * inch,
        bottomMargin=0.58 * inch,
        title=data["incident_name"],
        author="CyberVision",
        subject="Cybersecurity incident analysis",
    )
    story: list[Any] = [
        para("CYBERSECURITY INCIDENT REPORT", title),
        para(data["incident_name"], subtitle),
    ]
    metadata = [
        ("Incident ID", data["incident_id"]),
        ("Severity", data["severity"]),
        ("Created", _display_date(data)),
        ("Analysis engine", _engine_label(data)),
    ]
    for label, value in metadata:
        story.append(
            para(
                f"<b>{_reportlab_text(label)}:</b> {_reportlab_text(value)}",
                markup=True,
            )
        )
    story.append(Spacer(1, 4))

    summary_table = Table(
        [[para(data.get("summary") or "No executive summary was supplied.", body)]],
        colWidths=[6.5 * inch],
    )
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), palette["surface"]),
                ("BOX", (0, 0), (-1, -1), 0.75, palette["line"]),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.extend([summary_table, Spacer(1, 6)])

    mapped_steps = [step for step in data["steps"] if _known(step["mitre"]["technique_id"])]
    metrics = [
        ("Severity", data["severity"]),
        ("Pipeline quality", f'{data.get("confidence", 0)} / 100'),
        ("Attack steps", str(len(data["steps"]))),
        ("ATT&CK mapped", f"{len(mapped_steps)} / {len(data['steps'])}"),
    ]
    metric_cells = [
        [para(value, metric_value), para(label, metric_label)] for label, value in metrics
    ]
    metric_table = Table([metric_cells], colWidths=[1.625 * inch] * 4)
    metric_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOX", (0, 0), (-1, -1), 0.75, palette["line"]),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, palette["line"]),
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([metric_table, para("Attack sequence", heading)])

    step_headers = ["#", "Action", "Actor", "Target / Asset", "ATT&CK", "Detection"]
    step_rows = [[para(value, table_header) for value in step_headers]]
    for step in data["steps"]:
        mitre = (
            f'{_reportlab_text(step["mitre"]["technique_id"])}'
            f'<br/>{_reportlab_text(step["mitre"]["tactic"])}'
        )
        target = (
            f'{_reportlab_text(step["target"])}'
            f'<br/>{_reportlab_text(step["asset"])}'
        )
        step_rows.append(
            [
                para(step["order"], small),
                para(step["action"], small),
                para(step["actor"], small),
                para(target, small, markup=True),
                para(mitre, small, markup=True),
                para(step.get("detection"), small),
            ]
        )
    step_table = Table(
        step_rows,
        repeatRows=1,
        colWidths=[0.30 * inch, 1.25 * inch, 0.82 * inch, 1.00 * inch, 1.05 * inch, 2.08 * inch],
        hAlign="LEFT",
    )
    step_table.setStyle(_pdf_table_style(palette, data["severity"]))
    story.append(step_table)

    if graph_bytes:
        story.append(para("Attack diagram", heading))
        try:
            image = Image(io.BytesIO(graph_bytes))
            width, height = image.imageWidth, image.imageHeight
            scale = min((6.5 * inch) / width, (3.35 * inch) / height, 1.0)
            image.drawWidth = width * scale
            image.drawHeight = height * scale
            story.append(KeepTogether([image, Spacer(1, 4)]))
        except Exception:
            story.append(para("The supplied graph image could not be embedded.", small))

    story.append(para("MITRE ATT&CK mapping", heading))
    mitre_headers = ["Technique", "Tactic", "Detection", "Mitigation"]
    mitre_rows = [[para(value, table_header) for value in mitre_headers]]
    for step in data["steps"]:
        mitre_rows.append(
            [
                para(step["mitre"]["technique_id"], small),
                para(step["mitre"]["tactic"], small),
                para(step.get("detection"), small),
                para(step.get("mitigation"), small),
            ]
        )
    mitre_table = Table(
        mitre_rows,
        repeatRows=1,
        colWidths=[0.85 * inch, 1.15 * inch, 2.05 * inch, 2.45 * inch],
        hAlign="LEFT",
    )
    mitre_table.setStyle(_pdf_table_style(palette, data["severity"]))
    story.append(mitre_table)

    story.append(para("Entities and affected scope", heading))
    entities = data["entities"]
    story.extend(
        [
            para(
                f"<b>Actors:</b> {_reportlab_text(_join(entities.get('actors')))}",
                markup=True,
            ),
            para(
                f"<b>Targets:</b> {_reportlab_text(_join(entities.get('targets')))}",
                markup=True,
            ),
            para(
                f"<b>Assets:</b> {_reportlab_text(_join(entities.get('assets')))}",
                markup=True,
            ),
        ]
    )

    story.extend([PageBreak(), para("Recommended response", heading)])
    action_headers = ["Priority", "Response action", "Suggested owner", "Status"]
    action_rows = [[para(value, table_header) for value in action_headers]]
    response_owners = ("IR Lead", "IAM / Endpoint", "Threat Hunting")
    for index, action in enumerate(actions, 1):
        action_rows.append(
            [
                para("P1" if index <= 2 else "P2", small),
                para(action, small),
                para(response_owners[min(index - 1, len(response_owners) - 1)], small),
                para("Open", small),
            ]
        )
    action_table = Table(
        action_rows,
        repeatRows=1,
        colWidths=[0.65 * inch, 4.00 * inch, 1.10 * inch, 0.75 * inch],
        hAlign="LEFT",
    )
    action_table.setStyle(_pdf_table_style(palette, data["severity"]))
    story.append(action_table)

    metadata_block = data.get("metadata") or {}
    rag_metadata = metadata_block.get("rag") or {}
    orchestration = metadata_block.get("orchestration") or {}
    knowledge = metadata_block.get("knowledge") or {}
    provenance_rows = [
        ["Analysis engine", _engine_label(data)],
        [
            "Orchestration",
            str(orchestration.get("engine") or "native")
            + (
                f" · {orchestration.get('duration_ms')} ms"
                if orchestration.get("duration_ms") is not None
                else ""
            ),
        ],
        [
            "Semantic RAG",
            " · ".join(
                filter(
                    None,
                    (
                        str(rag_metadata.get("backend") or ""),
                        str(rag_metadata.get("embedding") or ""),
                        (
                            f"{rag_metadata.get('chunks')} chunks"
                            if rag_metadata.get("chunks") is not None
                            else ""
                        ),
                    ),
                )
            )
            or "Not available",
        ],
        [
            "Knowledge evidence",
            (
                f"{knowledge.get('matches', 0)} matches · "
                + ", ".join(knowledge.get("sources") or [])
            )
            if knowledge
            else "No multi-source evidence attached",
        ],
        [
            "Confidence method",
            str(
                (data.get("confidence_breakdown") or {}).get(
                    "methodology", "weighted_structure_quality_v1"
                )
            ),
        ],
    ]
    story.append(para("Analysis provenance", heading))
    provenance_table = Table(
        [
            [para(label, table_header), para(value, small)]
            for label, value in provenance_rows
        ],
        colWidths=[1.35 * inch, 5.15 * inch],
        hAlign="LEFT",
    )
    provenance_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.55, palette["line"]),
                ("BACKGROUND", (0, 0), (0, -1), palette["table"]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(provenance_table)
    def page_chrome(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        width, height = letter
        canvas.setStrokeColor(palette["line"])
        canvas.setLineWidth(0.6)
        canvas.line(inch, height - 0.48 * inch, width - inch, height - 0.48 * inch)
        canvas.setFont(font_bold, 7.5)
        canvas.setFillColor(palette["muted"])
        canvas.drawString(inch, height - 0.39 * inch, "CYBERVISION | INCIDENT REPORT")
        canvas.setFont(font_regular, 7.5)
        canvas.drawRightString(width - inch, 0.38 * inch, f"Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=page_chrome, onLaterPages=page_chrome)
    if not destination.exists() or destination.stat().st_size < 100:
        raise ReportGenerationError("ReportLab did not create a valid PDF.")
    return destination


def generate_docx(
    incident: dict[str, Any],
    output_path: str | os.PathLike[str] | None = None,
    *,
    graph_image: bytes | bytearray | io.BytesIO | str | os.PathLike[str] | None = None,
    recommendations: Iterable[str] | None = None,
) -> Path:
    """Create a Word report using the standard_business_brief preset.

    The first page uses the memo_masthead pattern.  All page, paragraph and
    table geometry is explicit so the document is deterministic across Word and
    LibreOffice.
    """
    data = _canonical_incident(incident)
    destination = _prepare_output(data, output_path, ".docx")
    graph_bytes = _read_binary(graph_image)
    actions = _recommendations(data, recommendations)

    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ReportGenerationError(
            "DOCX export requires python-docx>=1.1."
        ) from exc

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    _configure_docx_styles(document)
    _configure_docx_header_footer(document, data)

    title_paragraph = document.add_paragraph(style="CVI Title")
    title_paragraph.add_run("CYBERSECURITY INCIDENT REPORT")
    subtitle_paragraph = document.add_paragraph(style="CVI Subtitle")
    subtitle_paragraph.add_run(data["incident_name"])

    metadata = [
        ("To", "SOC / Incident Response Team"),
        ("From", "CyberVision Analysis Pipeline"),
        ("Date", _display_date(data)),
        ("Incident ID", data["incident_id"]),
        ("Status", f'{data["severity"]} severity - review and response required'),
    ]
    for index, (label, value) in enumerate(metadata):
        paragraph = document.add_paragraph(style="CVI Metadata")
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)
        if index == len(metadata) - 1:
            _set_paragraph_bottom_border(paragraph, color="0B2545", size=12, space=7)

    document.add_heading("Executive summary", level=1)
    summary_table = document.add_table(rows=1, cols=1)
    summary_table.style = "Table Grid"
    _set_table_geometry(summary_table, [9360])
    _shade_cell(summary_table.cell(0, 0), "F4F6F9")
    _set_cell_text(summary_table.cell(0, 0), data.get("summary") or "No summary supplied.")

    mapped_steps = [step for step in data["steps"] if _known(step["mitre"]["technique_id"])]
    document.add_heading("Analysis snapshot", level=1)
    metrics = [
        ("Severity", data["severity"]),
        ("Pipeline quality", f'{data.get("confidence", 0)} / 100'),
        ("Attack steps", str(len(data["steps"]))),
        ("ATT&CK mapped", f"{len(mapped_steps)} / {len(data['steps'])}"),
    ]
    metric_table = document.add_table(rows=2, cols=4)
    metric_table.style = "Table Grid"
    _set_table_geometry(metric_table, [2340, 2340, 2340, 2340])
    for column, (label, value) in enumerate(metrics):
        _set_cell_text(metric_table.cell(0, column), value, bold=True, center=True)
        _set_cell_text(metric_table.cell(1, column), label, center=True, muted=True)
        _shade_cell(metric_table.cell(1, column), "F2F4F7")

    document.add_heading("Attack sequence", level=1)
    step_widths = [520, 2150, 1320, 1500, 1660, 2210]
    step_headers = ["#", "Action", "Actor", "Target / Asset", "ATT&CK", "Detection"]
    step_table = document.add_table(rows=1, cols=len(step_headers))
    step_table.style = "Table Grid"
    _set_table_geometry(step_table, step_widths)
    _mark_repeat_header(step_table.rows[0])
    for index, header in enumerate(step_headers):
        _set_cell_text(step_table.cell(0, index), header, bold=True)
        _shade_cell(step_table.cell(0, index), "F2F4F7")
    for step in data["steps"]:
        cells = step_table.add_row().cells
        values = [
            step["order"],
            step["action"],
            step["actor"],
            f'{step["target"]}\n{step["asset"]}',
            f'{step["mitre"]["technique_id"]}\n{step["mitre"]["tactic"]}',
            step.get("detection"),
        ]
        for index, value in enumerate(values):
            _set_cell_text(cells[index], value, center=index == 0)
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _apply_row_widths(cells, step_widths)

    if graph_bytes:
        document.add_heading("Attack diagram", level=1)
        stream = io.BytesIO(graph_bytes)
        graph_paragraph = document.add_paragraph()
        graph_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        graph_paragraph.add_run().add_picture(stream, width=Inches(6.35))
        caption = document.add_paragraph("Figure 1. Generated attack flow", style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("MITRE ATT&CK mapping", level=1)
    mitre_widths = [1200, 1650, 2960, 3550]
    mitre_headers = ["Technique", "Tactic", "Detection", "Mitigation"]
    mitre_table = document.add_table(rows=1, cols=4)
    mitre_table.style = "Table Grid"
    _set_table_geometry(mitre_table, mitre_widths)
    _mark_repeat_header(mitre_table.rows[0])
    for index, header in enumerate(mitre_headers):
        _set_cell_text(mitre_table.cell(0, index), header, bold=True)
        _shade_cell(mitre_table.cell(0, index), "F2F4F7")
    for step in data["steps"]:
        cells = mitre_table.add_row().cells
        values = [
            step["mitre"]["technique_id"],
            step["mitre"]["tactic"],
            step.get("detection"),
            step.get("mitigation"),
        ]
        for index, value in enumerate(values):
            _set_cell_text(cells[index], value)
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _apply_row_widths(cells, mitre_widths)

    document.add_heading("Entities and affected scope", level=1)
    entities = data["entities"]
    for label, values in (
        ("Actors", entities.get("actors")),
        ("Targets", entities.get("targets")),
        ("Assets", entities.get("assets")),
    ):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(_join(values))

    document.add_heading("Recommended response", level=1)
    response_table = document.add_table(rows=1, cols=2)
    response_table.style = "Table Grid"
    response_widths = [1400, 7960]
    _set_table_geometry(response_table, response_widths)
    for index, header in enumerate(("Priority", "Action")):
        _set_cell_text(response_table.cell(0, index), header, bold=True)
        _shade_cell(response_table.cell(0, index), "F2F4F7")
    _mark_repeat_header(response_table.rows[0])
    for index, action in enumerate(actions, 1):
        cells = response_table.add_row().cells
        _set_cell_text(cells[0], "P1" if index <= 2 else "P2", bold=True, center=True)
        _set_cell_text(cells[1], action)
        _apply_row_widths(cells, response_widths)

    document.add_heading("Methodology note", level=1)
    document.add_paragraph(
        "The confidence value is an explainable pipeline-quality score derived from "
        "structured-field completeness, model source reliability and ATT&CK RAG "
        "coverage. It is not a probability that the incident occurred."
    )
    properties = document.core_properties
    properties.title = data["incident_name"]
    properties.subject = "Cybersecurity incident analysis"
    properties.author = "CyberVision"
    properties.keywords = "cybersecurity, MITRE ATT&CK, incident response"
    properties.comments = (
        "Generated with standard_business_brief and memo_masthead design contracts."
    )
    document.save(destination)
    if not destination.exists() or destination.stat().st_size < 1000:
        raise ReportGenerationError("python-docx did not create a valid DOCX.")
    return destination


def generate_pptx(
    incident: dict[str, Any],
    output_path: str | os.PathLike[str] | None = None,
    *,
    graph_image: bytes | bytearray | io.BytesIO | str | os.PathLike[str] | None = None,
    recommendations: Iterable[str] | None = None,
    node_binary: str | os.PathLike[str] | None = None,
    artifact_tool_node_modules: str | os.PathLike[str] | None = None,
    qa_directory: str | os.PathLike[str] | None = None,
) -> Path:
    """Create an editable PPTX exclusively through @oai/artifact-tool.

    The Node runtime and package directory are discovered without installing
    dependencies.  Deployments may pin them with ``CVI_NODE_BINARY`` and
    ``ARTIFACT_TOOL_NODE_MODULES``.
    """
    data = _canonical_incident(incident)
    destination = _prepare_output(data, output_path, ".pptx")
    node = _find_node_binary(node_binary)
    modules = _find_artifact_tool_modules(artifact_tool_node_modules)
    if node is None:
        raise ReportGenerationError(
            "PPTX export requires Node.js. Set CVI_NODE_BINARY to node/node.exe."
        )
    if modules is None:
        raise ReportGenerationError(
            "PPTX export requires @oai/artifact-tool. Set "
            "ARTIFACT_TOOL_NODE_MODULES to the node_modules directory containing "
            "@oai/artifact-tool."
        )
    if not PPTX_SCRIPT.exists():
        raise ReportGenerationError(f"PPTX generator script is missing: {PPTX_SCRIPT}")

    graph_bytes = _read_binary(graph_image)
    payload = {
        "incident": data,
        "recommendations": _recommendations(data, recommendations),
        "graph_image_base64": base64.b64encode(graph_bytes).decode("ascii") if graph_bytes else None,
        "graph_image_content_type": _detect_image_content_type(graph_bytes),
    }
    temp_root = Path(tempfile.mkdtemp(prefix="cvi-pptx-"))
    payload_path = temp_root / "incident-report.json"
    temporary_qa_directory = qa_directory is None
    scratch = (
        Path(qa_directory).resolve()
        if qa_directory
        else Path(
            tempfile.mkdtemp(prefix=f'cvi-pptx-qa-{_slug(data["incident_id"])}-')
        )
    )
    scratch.mkdir(parents=True, exist_ok=True)
    payload_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    command = [
        str(node),
        str(PPTX_SCRIPT),
        str(payload_path),
        str(destination),
        str(modules),
        str(scratch),
    ]
    try:
        completed = subprocess.run(
            command,
            cwd=str(PROJECT_ROOT),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise ReportGenerationError(f"PPTX exporter could not run: {exc}") from exc
    finally:
        try:
            payload_path.unlink(missing_ok=True)
            shutil.rmtree(temp_root, ignore_errors=True)
            if temporary_qa_directory:
                shutil.rmtree(scratch, ignore_errors=True)
        except OSError:
            pass
    if completed.returncode != 0:
        streams = [
            value.strip()
            for value in (completed.stderr, completed.stdout)
            if value and value.strip()
        ]
        details = "\n".join(streams) or "Unknown artifact-tool error"
        raise ReportGenerationError(f"@oai/artifact-tool PPTX export failed: {details}")
    if not destination.exists() or destination.stat().st_size < 1000:
        raise ReportGenerationError("@oai/artifact-tool did not create a valid PPTX.")
    return destination


def generate_report(
    report_format: str,
    incident: dict[str, Any],
    output_path: str | os.PathLike[str] | None = None,
    **kwargs: Any,
) -> Path:
    """Dispatch one of the supported server-side report exporters."""
    normalized = str(report_format).strip().lower().lstrip(".")
    exporters = {"pdf": generate_pdf, "docx": generate_docx, "pptx": generate_pptx}
    if normalized not in exporters:
        raise ValueError("report_format must be one of: pdf, docx, pptx")
    return exporters[normalized](incident, output_path, **kwargs)


def report_capabilities() -> dict[str, Any]:
    """Return truthful runtime capability metadata for a health/config API."""
    node = _find_node_binary()
    artifact_modules = _find_artifact_tool_modules()
    try:
        import reportlab  # noqa: F401

        pdf_ready = True
    except ImportError:
        pdf_ready = False
    try:
        import docx  # noqa: F401

        docx_ready = True
    except ImportError:
        docx_ready = False
    return {
        "pdf": {"ready": pdf_ready, "engine": "reportlab"},
        "docx": {
            "ready": docx_ready,
            "engine": "python-docx",
            "preset": "standard_business_brief",
            "header": "memo_masthead",
        },
        "pptx": {
            "ready": bool(node and artifact_modules),
            "engine": "@oai/artifact-tool",
            "node": str(node) if node else None,
            "node_modules": str(artifact_modules) if artifact_modules else None,
        },
    }


def _canonical_incident(value: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise ValueError("incident must be an object")
    candidate = value.get("structured_json", value)
    if not isinstance(candidate, dict):
        raise ValueError("structured_json must be an object")
    return validate_structured_incident(candidate)


def _prepare_output(
    incident: dict[str, Any],
    output_path: str | os.PathLike[str] | None,
    suffix: str,
) -> Path:
    if output_path is None:
        destination = DEFAULT_REPORT_DIR / f'{_slug(incident["incident_id"])}-report{suffix}'
    else:
        destination = Path(output_path).expanduser()
        if destination.suffix.lower() != suffix:
            destination = destination.with_suffix(suffix)
    destination = destination.resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    return destination


def _register_pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    regular_candidates = [
        os.getenv("REPORT_FONT_REGULAR"),
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    bold_candidates = [
        os.getenv("REPORT_FONT_BOLD"),
        r"C:\Windows\Fonts\arialbd.ttf",
        r"C:\Windows\Fonts\calibrib.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf",
    ]
    regular = _first_file(regular_candidates)
    bold = _first_file(bold_candidates)
    if regular:
        try:
            pdfmetrics.registerFont(TTFont("CVIText", str(regular)))
            pdfmetrics.registerFont(TTFont("CVIText-Bold", str(bold or regular)))
            return "CVIText", "CVIText-Bold"
        except Exception:
            pass
    return "Helvetica", "Helvetica-Bold"


def _pdf_table_style(palette: dict[str, Any], severity: str) -> Any:
    from reportlab.platypus import TableStyle

    return TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), palette["table"]),
            ("TEXTCOLOR", (0, 0), (-1, 0), palette["navy"]),
            ("GRID", (0, 0), (-1, -1), 0.55, palette["line"]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LINEBEFORE", (0, 0), (0, -1), 1.6, palette["severity"]),
        ]
    )


def _configure_docx_styles(document: Any) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor

    normal = document.styles["Normal"]
    _set_style_font(normal, "Calibri", 11, RGBColor(0x0B, 0x25, 0x45))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ):
        style = document.styles[name]
        _set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = document.styles.add_style("CVI Title", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(title, "Calibri", 23, RGBColor(0, 0, 0), bold=True)
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)

    subtitle = document.styles.add_style("CVI Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(subtitle, "Calibri", 14, RGBColor(0x37, 0x37, 0x37))
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)

    metadata = document.styles.add_style("CVI Metadata", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(metadata, "Calibri", 11, RGBColor(0, 0, 0))
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing = 1.0

    caption = document.styles["Caption"]
    _set_style_font(caption, "Calibri", 9, RGBColor(0x5E, 0x6C, 0x7B), italic=True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)


def _configure_docx_header_footer(document: Any, data: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("CYBERVISION  |  INCIDENT REPORT")
    left.bold = True
    _set_run_font(left, "Calibri", 8, "5E6C7B")
    paragraph.add_run("\t")
    right = paragraph.add_run(data["incident_id"])
    _set_run_font(right, "Calibri", 8, "5E6C7B")
    _set_right_tab(paragraph, 6.5)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ""
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_paragraph.add_run("Page ")
    _set_run_font(run, "Calibri", 8, "5E6C7B")
    _append_field(footer_paragraph, "PAGE")
    footer_paragraph.add_run(" of ")
    _append_field(footer_paragraph, "NUMPAGES")


def _set_style_font(
    style: Any,
    name: str,
    size: float,
    color: Any,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def _set_run_font(run: Any, name: str, size: float, color: str) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _set_table_geometry(
    table: Any,
    widths_dxa: list[int],
    *,
    indent_dxa: int = 120,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if sum(widths_dxa) != 9360:
        raise ValueError("DOCX table widths must sum to 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    def child(tag: str) -> Any:
        element = tbl_pr.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tbl_pr.append(element)
        return element

    tbl_w = child("tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = child("tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_layout = child("tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")

    margins = child("tblCellMar")
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        side_element = margins.find(qn(f"w:{side}"))
        if side_element is None:
            side_element = OxmlElement(f"w:{side}")
            margins.append(side_element)
        side_element.set(qn("w:type"), "dxa")
        side_element.set(qn("w:w"), str(value))

    grid = table._tbl.tblGrid
    for element in list(grid):
        grid.remove(element)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        _apply_row_widths(row.cells, widths_dxa)


def _apply_row_widths(cells: Any, widths_dxa: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Twips

    for index, (cell, width) in enumerate(zip(cells, widths_dxa)):
        cell.width = Twips(width)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(widths_dxa[index]))


def _set_cell_text(
    cell: Any,
    value: Any,
    *,
    bold: bool = False,
    center: bool = False,
    muted: bool = False,
) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    parts = str(value if value not in (None, "") else "Unknown").splitlines() or ["Unknown"]
    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        _set_run_font(run, "Calibri", 8.5, "5E6C7B" if muted else "0B2545")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _mark_repeat_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def _set_paragraph_bottom_border(
    paragraph: Any,
    *,
    color: str,
    size: int,
    space: int,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def _set_right_tab(paragraph: Any, inches: float) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches

    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(inches), WD_TAB_ALIGNMENT.RIGHT
    )


def _append_field(paragraph: Any, field_name: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    _set_run_font(run, "Calibri", 8, "5E6C7B")


def _find_node_binary(
    explicit: str | os.PathLike[str] | None = None,
) -> Path | None:
    candidates: list[Path] = []
    if explicit:
        candidates.append(Path(explicit))
    if os.getenv("CVI_NODE_BINARY"):
        candidates.append(Path(os.environ["CVI_NODE_BINARY"]))
    local = shutil.which("node")
    if local:
        candidates.append(Path(local))
    candidates.extend(
        [
            PROJECT_ROOT / ".tools" / "node" / "node.exe",
            PROJECT_ROOT / ".tools" / "node" / "bin" / "node.exe",
        ]
    )
    candidates.extend(
        sorted(PROJECT_ROOT.glob(".tools/node*/**/node.exe"), reverse=True)
    )
    local_app_data = os.getenv("LOCALAPPDATA")
    if local_app_data:
        runtime_root = Path(local_app_data) / "OpenAI" / "Codex" / "runtimes" / "cua_node"
        if runtime_root.exists():
            candidates.extend(sorted(runtime_root.glob("*/bin/node.exe"), reverse=True))
    return _first_file(candidates)


def _find_artifact_tool_modules(
    explicit: str | os.PathLike[str] | None = None,
) -> Path | None:
    candidates: list[Path] = []
    if explicit:
        candidates.append(Path(explicit))
    if os.getenv("ARTIFACT_TOOL_NODE_MODULES"):
        candidates.append(Path(os.environ["ARTIFACT_TOOL_NODE_MODULES"]))
    candidates.extend(
        [
            PROJECT_ROOT / "node_modules",
            PROJECT_ROOT / ".tools" / "artifact-tool" / "node_modules",
        ]
    )
    # Codex primary-runtime packages are installed outside the workspace.  The
    # location is intentionally discovered at runtime instead of copied into
    # the repository, so the Flask service can use the supported private
    # artifact-tool bundle without vendoring it.
    home_candidates = {
        Path.home(),
        Path(os.environ["HOME"]).expanduser() if os.getenv("HOME") else None,
        (
            Path(os.environ["USERPROFILE"]).expanduser()
            if os.getenv("USERPROFILE")
            else None
        ),
    }
    for home in (candidate for candidate in home_candidates if candidate):
        candidates.append(
            home
            / ".cache"
            / "codex-runtimes"
            / "codex-primary-runtime"
            / "dependencies"
            / "node"
            / "node_modules"
        )
    local_app_data = os.getenv("LOCALAPPDATA")
    if local_app_data:
        runtime_root = Path(local_app_data) / "OpenAI" / "Codex" / "runtimes"
        if runtime_root.exists():
            candidates.extend(runtime_root.glob("**/bin/node_modules"))
    for candidate in candidates:
        path = candidate.expanduser()
        if (path / "@oai" / "artifact-tool" / "package.json").is_file():
            return path.resolve()
    return None


def _recommendations(
    data: dict[str, Any],
    supplied: Iterable[str] | None,
) -> list[str]:
    if supplied:
        unique = [_clean_text(item) for item in supplied]
        return list(dict.fromkeys(item for item in unique if item))[:8]

    actions: list[str] = [
        "Isolate affected endpoints and suspend compromised accounts while preserving forensic evidence.",
        "Block validated indicators across email, DNS, proxy, EDR and firewall controls.",
    ]
    tactics = {step["mitre"]["tactic"] for step in data["steps"]}
    if "Credential Access" in tactics:
        actions.append(
            "Revoke active sessions, rotate exposed credentials and enforce phishing-resistant MFA."
        )
    if "Command and Control" in tactics or "Command and Control" in {
        value.title() for value in tactics
    }:
        actions.append(
            "Hunt for related command-and-control traffic and quarantine hosts with matching telemetry."
        )
    detections = [
        _clean_text(step.get("detection"))
        for step in data["steps"]
        if _known(step.get("detection"))
    ]
    if detections:
        actions.append(
            "Operationalize the mapped detections: " + "; ".join(dict.fromkeys(detections))[:650]
        )
    technique_ids = [
        step["mitre"]["technique_id"]
        for step in data["steps"]
        if _known(step["mitre"]["technique_id"])
    ]
    if technique_ids:
        actions.append(
            "Run environment-wide threat hunting for ATT&CK techniques "
            + ", ".join(dict.fromkeys(technique_ids))
            + "."
        )
    return actions[:8]


def _read_binary(
    source: bytes | bytearray | io.BytesIO | str | os.PathLike[str] | None,
) -> bytes | None:
    if source is None:
        return None
    if isinstance(source, bytes):
        return source
    if isinstance(source, bytearray):
        return bytes(source)
    if hasattr(source, "getvalue"):
        return bytes(source.getvalue())
    path = Path(source)
    return path.read_bytes()


def _detect_image_content_type(content: bytes | None) -> str | None:
    if not content:
        return None
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if content.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if content.lstrip().startswith(b"<svg"):
        return "image/svg+xml"
    return "application/octet-stream"


def _display_date(data: dict[str, Any]) -> str:
    raw = data.get("metadata", {}).get("created_at")
    if raw:
        try:
            parsed = datetime.fromisoformat(str(raw).replace("Z", "+00:00"))
            return parsed.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        except ValueError:
            return str(raw)
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _engine_label(data: dict[str, Any]) -> str:
    metadata = data.get("metadata", {})
    model = _clean_text(metadata.get("model")) or "Unknown"
    provider = _clean_text(metadata.get("provider")) or "Unknown"
    return f"{model} / {provider}"


def _severity_color(severity: str) -> str:
    return {
        "Critical": "#B42318",
        "High": "#D97706",
        "Medium": "#CA8A04",
        "Low": "#15803D",
        "Unknown": "#607086",
    }.get(severity, "#607086")


def _reportlab_text(value: Any) -> str:
    from xml.sax.saxutils import escape

    text = _clean_text(value) or "Unknown"
    return escape(text).replace("\n", "<br/>")


def _join(values: Any) -> str:
    if not values:
        return "Unknown"
    if isinstance(values, str):
        return values
    return ", ".join(str(value) for value in values if _known(value)) or "Unknown"


def _known(value: Any) -> bool:
    return _clean_text(value).lower() not in {"", "unknown", "none", "null", "n/a"}


def _clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _slug(value: Any) -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "-", str(value)).strip("-._")
    return slug or "incident"


def _first_file(candidates: Iterable[str | os.PathLike[str] | None]) -> Path | None:
    for candidate in candidates:
        if not candidate:
            continue
        path = Path(candidate).expanduser()
        if path.is_file():
            return path.resolve()
    return None
